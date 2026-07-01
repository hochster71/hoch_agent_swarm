import os
from docx import Document

class DocFactory:
    def create_brief(self, title: str, sections: list, filepath: str) -> str:
        doc = Document()
        doc.add_heading(title, 0)
        
        for heading, body in sections:
            doc.add_heading(heading, level=1)
            doc.add_paragraph(body)
            
        os.makedirs(os.path.dirname(filepath), exist_ok=True)
        doc.save(filepath)
        return filepath
